import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x99) # Teal/Blue
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h

def create_briefing_document():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # Title & Metadata Banner
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_title = title_p.add_run("Shared Artlist Asset Library & Automated Relay")
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    run_sub = sub_p.add_run("Executive Briefing, System Architecture & Operational Risk Analysis (v1.3)")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Callout / Meta Box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.cell(0, 0)
    set_cell_background(c, "F0F4F8")
    set_cell_margins(c, top=140, bottom=140, left=180, right=180)
    cp = c.paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    cp.paragraph_format.line_spacing = 1.15
    r_meta = cp.add_run("Target Audience: ")
    r_meta.bold = True
    cp.add_run("Team Lead, Video Editors, Stakeholders & Developers\n")
    r_meta2 = cp.add_run("Core Purpose: ")
    r_meta2.bold = True
    cp.add_run("Eliminate redundant licensing, build a permanent shared org music library, automate asset delivery to Google Drive, and preserve individual daily download quotas.")

    doc.add_paragraph() # spacing

    # 1. Executive Summary
    add_styled_heading(doc, "1. Executive Summary & Why We Are Building This", level=1)
    
    p = doc.add_paragraph()
    p.add_run("Our video editing team currently operates with 5 independent Artlist seats. However, because each editor works in isolation, we face five daily workflow problems:")
    
    bullets = [
        ("Redundant Licensing & Downloads: ", "The exact same track often gets downloaded 3 to 5 times by different editors for different projects, wasting time and quota."),
        ("No Shared Team Library: ", "Files land in local Downloads folders. There is no centralized place to check before searching Artlist."),
        ("Manual File Sharing Friction: ", "Sharing a track currently requires manual download, navigating to Google Drive, uploading, and notifying colleagues—which is often skipped under deadline."),
        ("Personal Quota Drain: ", "Each editor's 40-track daily allowance is consumed by re-downloading music the organization already owns."),
        ("No Centralized Audit Trail: ", "The team lacks a unified record of what has been licensed, when, and for which project.")
    ]
    for b_title, b_desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r_bt = bp.add_run(b_title)
        r_bt.bold = True
        bp.add_run(b_desc)

    p_sol = doc.add_paragraph()
    p_sol.paragraph_format.space_before = Pt(6)
    p_sol.add_run("The Solution: ").bold = True
    p_sol.add_run("An automated, dedicated relay node (Mac or Windows) linked to a designated Artlist relay account. Editors simply drop a link into Google Chat or the local web portal. The system checks a local cache (returning existing tracks in <1s at zero quota), or automatically downloads the high-quality WAV via an authenticated browser session, atomically moves it to a shared Google Drive library, and replies in-thread.")

    # 2. How the System Works (Architecture)
    add_styled_heading(doc, "2. System Architecture & End-to-End Flow", level=1)

    steps = [
        ("1. Ingestion Layer (Google Chat / Web Portal): ", "Editors paste an Artlist link into Google Chat (#artlist-library). Using Cloud Pub/Sub pull subscriptions, the relay receives the message securely with ZERO inbound ports or firewall exposure. Identity is verified via Google Workspace."),
        ("2. Deduplication & Cache Engine (FastAPI + SQLite WAL): ", "The server extracts the track ID and checks SQLite. If the track already exists in the Shared Drive, it returns the file path in <3s with 0 quota spent. Over time, >40% of requests become instant cache hits."),
        ("3. Execution Engine (Native Chrome + MV3 Extension): ", "For new tracks, an unpacked/force-installed Chrome Extension on the dedicated node triggers the download inside an authenticated human session using human-paced delays (4s–9s)."),
        ("4. Two-Stage File Delivery (Staging -> Shared Drive): ", "Chrome downloads the file to an isolated local staging directory. Once byte-size and RIFF/WAVE headers are validated, the server performs an atomic move (os.replace) into the Google Drive synced folder, ensuring editors never see partial or .crdownload files.")
    ]
    for s_title, s_desc in steps:
        sp = doc.add_paragraph(style='List Number')
        sp.paragraph_format.space_after = Pt(4)
        r_st = sp.add_run(s_title)
        r_st.bold = True
        sp.add_run(s_desc)

    # 3. Key Risks & Mitigations
    add_styled_heading(doc, "3. Key Operational & Technical Risks (What Could Go Wrong & Solutions)", level=1)
    
    intro_r = doc.add_paragraph()
    intro_r.add_run("Below is the curated risk assessment covering technical, operational, and organizational factors:")

    # Risk Table
    risk_table = doc.add_table(rows=1, cols=4)
    risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = risk_table.rows[0].cells
    headers = ["Risk Area", "Potential Failure Mode", "Impact", "Mitigation Strategy"]
    widths = [Inches(1.3), Inches(2.0), Inches(0.8), Inches(2.7)]

    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "1B365D")
        set_cell_margins(hdr_cells[idx], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)

    risk_data = [
        ("1. Editor Adoption & Habit", 
         "Editors bypass the bot and click 'Download' directly on their personal Artlist seats out of habit.", 
         "High", 
         "Implement 1-click Raycast/Alfred hotkey or browser bookmarklet. Emphasize that library tracks preserve personal quota."),
        
        ("2. IT / Workspace Approval Delay", 
         "Google Workspace Admin takes days/weeks to approve the Google Chat App & GCP project.", 
         "Medium", 
         "Deploy the standalone Web Portal (http://node.local:5000) on Day 1 as an immediate functional bridge while Chat app is reviewed."),
        
        ("3. Account Ban / Flagging", 
         "Artlist detects automation and flags or locks the account.", 
         "Medium", 
         "Blast Radius Contained: Relay runs on 1 dedicated account. The 5 editors' seats remain completely unaffected. Pacing delays & 35-cap safety stop enforced."),
        
        ("4. MV3 Service Worker Sleep", 
         "Chrome Manifest V3 terminates background workers after 30s idle, freezing the poll loop.", 
         "High", 
         "Use chrome.alarms (30s period) to reliably wake the worker, plus active tab/download listeners to keep it alive during jobs."),
         
        ("5. Artlist Frontend Changes", 
         "Artlist changes DOM classes, layouts, or adds promotional/cookie popups.", 
         "Medium", 
         "Centralized fallback selector dictionary + automated Circuit Breaker that auto-pauses the queue and alerts lead after 3 failures."),
         
        ("6. Storage Quota Exhaustion", 
         "Accumulating 60MB WAV files fills local SSD or Google Drive cloud storage.", 
         "Medium", 
         "Add disk headroom telemetry. Alert team lead at 80% Drive capacity; hard pause at 95%. Provide cold-track reporting."),
         
        ("7. Multi-Variant / Stems Confusion", 
         "Bot downloads Main vocal track when the editor needed Instrumental or Stems.", 
         "Medium", 
         "Explicit 'variant' parameter in schema; composite (track_id, variant) key in cache; strict selector mapping for track versions.")
    ]

    for item in risk_data:
        row_cells = risk_table.add_row().cells
        for col_idx, text in enumerate(item):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            for run in p.runs:
                run.font.size = Pt(9)
                if col_idx == 2: # Impact
                    if text == "High":
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                    else:
                        run.font.color.rgb = RGBColor(0xDD, 0x66, 0x00)
                elif col_idx == 0:
                    run.font.bold = True

    doc.add_paragraph() # spacing

    # 4. Phase 0 Checklist & Decisions
    add_styled_heading(doc, "4. Key Decisions & Phase 0 Action Items", level=1)

    p_d = doc.add_paragraph()
    p_d.add_run("Before full implementation begins, the following 4 decisions/actions must be executed:")
    
    decisions = [
        ("Submit Google Chat App Request (Longest Lead Time): ", "Set up GCP project, Pub/Sub topic, and submit Chat App manifest to the Google Workspace Super Admin for org-wide allowlisting."),
        ("Designate Node Hardware: ", "Confirm whether the always-on node is the Mac mini or a dedicated Windows 11 laptop (mains-powered, wired ethernet, sleep/hibernation disabled)."),
        ("Select Google Drive Sync Model: ", "Option A (Local Mirrored Folder) vs. Option B (Shared Drive Streamed). Option A is recommended for immediate atomic local moves."),
        ("Relay Account Allocation: ", "Confirm whether the relay runs on a dedicated 6th seat or a repurposed team seat.")
    ]
    for d_title, d_desc in decisions:
        dp = doc.add_paragraph(style='List Bullet')
        dp.paragraph_format.space_after = Pt(3)
        r_dt = dp.add_run(d_title)
        r_dt.bold = True
        dp.add_run(d_desc)

    # 5. Executive Talking Points
    add_styled_heading(doc, "5. Executive Talking Points for Stakeholders", level=1)
    
    talk_points = [
        "\"This system is purely additive.\" — It does not restrict or replace anyone's existing Artlist login. Editors can still download manually anytime, but now have a 1-click way to build a shared team library.",
        "\"Cache hits cost zero quota.\" — Once a track is requested, every subsequent use by any editor is instant (<1s) and saves team download allowance.",
        "\"Zero security vulnerabilities.\" — Using Cloud Pub/Sub pull architecture means zero open ports, zero tunnels, and no firewall changes on the office network.",
        "\"No corrupt partial files.\" — Staging on local disk + atomic handoff ensures files only appear in Google Drive once fully downloaded and verified.",
        "\"Rapid 3-day implementation.\" — Phase 1 (Core engine), Phase 2 (Extension), Phase 3 (Chat Bot / Web UI), and Phase 4 (Team rollout)."
    ]
    for tp in talk_points:
        tpp = doc.add_paragraph(style='List Bullet')
        tpp.paragraph_format.space_after = Pt(3)
        tpp.add_run(tp)

    # Save
    output_path = r"d:\Code\Artlist\Artlist-Relay-Architecture-and-Risk-Briefing.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    create_briefing_document()
